from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# -----------------------------
# Adds a clickable hyperlink to a paragraph
# -----------------------------
def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    # Create hyperlink XML element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run element inside the hyperlink
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')  # Run properties

    # Style as hyperlink (blue and underlined)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)

    # Set the display text of the hyperlink
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

# -----------------------------
# Adds visible borders to a table cell
# -----------------------------
def set_cell_borders(cell, color="A4A4A4"):
    cell._element.get_or_add_tcPr().append(parse_xml(f'''
        <w:tcBorders {nsdecls('w')}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    '''))

# -----------------------------
# Adds internal cell padding (margins) for readability
# -----------------------------
def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(value))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

# -----------------------------
# Main function to write a new Word document based on parsed content
# -----------------------------
def write_new_doc(parsed_content, output_path, template_path="templates/Article-Template.docx"):
    # Load the template document or create a blank one
    doc = Document(template_path) if template_path else Document()

    # Remove any placeholder paragraphs from the template
    if doc.paragraphs:
        for _ in range(len(doc.paragraphs)):
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)

    # Insert the title using Heading 1 style
    doc.add_paragraph(parsed_content["title"], style="Heading 1")

    # Loop through all content blocks parsed from the original document
    for block in parsed_content["structured_body"]:
        block_type = block["type"]

        # Insert headings
        if block_type == "heading2":
            doc.add_paragraph(block["text"], style="Heading 2")

        elif block_type == "heading3":
            doc.add_paragraph(block["text"], style="Heading 3")

        # Insert normal paragraphs
        elif block_type == "paragraph":
            para = doc.add_paragraph(style="Body Text")
            for run in block.get("runs", []):
                if run["hyperlink"]:
                    add_hyperlink(para, run["hyperlink"], run["text"])
                else:
                    r = para.add_run(run["text"])
                    if run["bold"]: r.bold = True
                    if run["italic"]: r.italic = True
                    if run["underline"]: r.underline = True

        # Insert list-style paragraphs (unordered or ordered—same handling here)
        elif block_type == "list_item":
            para = doc.add_paragraph(style="List Paragraph")
            for run in block.get("runs", []):
                r = para.add_run(run["text"])
                if run["bold"]: r.bold = True
                if run["italic"]: r.italic = True
                if run["underline"]: r.underline = True

        # Insert tables
        elif block_type == "table":
            table_data = block["rows"]
            num_cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=0, cols=num_cols)

            for row_cells in table_data:
                row = table.add_row().cells
                for col_idx, cell_runs in enumerate(row_cells):
                    cell = row[col_idx]
                    set_cell_borders(cell)        # Add light grey borders
                    set_cell_margins(cell)        # Add readable spacing inside
                    para = cell.paragraphs[0]     # Add text to first paragraph in cell
                    for run in cell_runs:
                        r = para.add_run(run["text"])
                        if run["bold"]: r.bold = True
                        if run["italic"]: r.italic = True
                        if run["underline"]: r.underline = True

            doc.add_paragraph("")  # Add spacing after the table

    # Save the finished document
    doc.save(output_path)