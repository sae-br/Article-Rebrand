from docx2python import docx2python  # For textboxes (title extraction)
from docx import Document
from docx.oxml.ns import qn
import json

# -----------------------------
# Load known authors from a file
# -----------------------------
def load_known_authors(filepath="authors.json"):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)

# -----------------------------
# Load raw .docx file using docx2python
# -----------------------------
def load_docx(filepath):
    return docx2python(filepath)

# -----------------------------
# Check if a paragraph is a list item (uses XML structure)
# -----------------------------
def is_list_paragraph(paragraph):
    return bool(paragraph._element.xpath('./w:pPr/w:numPr'))

# -----------------------------
# Extract all textboxes (used for title/byline detection)
# -----------------------------
def extract_textboxes(doc_structure):
    textboxes = []
    for section in [doc_structure.body] + ([doc_structure.header] if hasattr(doc_structure, 'header') else []):
        for table in section:
            for row in table:
                for cell in row:
                    text = " ".join(cell).strip()
                    if text:
                        textboxes.append(text)
    return textboxes

# -----------------------------
# Extract the most likely title from the textboxes
# -----------------------------
def choose_title_box(textboxes):
    for text in textboxes:
        if "an article by" in text.lower():
            continue
        if len(text.split()) < 15:
            return clean_title(text)
    return ""

def clean_title(text):
    return " ".join(text.split())

# -----------------------------
# Identify known author name from first few paragraphs
# -----------------------------
def detect_author(structured_body, known_authors, max_paragraphs=3):
    paragraphs = [block["text"] for block in structured_body if block["type"] == "paragraph"]
    search_text = " ".join(paragraphs[:max_paragraphs]).lower()
    for author in known_authors:
        if author.lower() in search_text:
            return author
    return ""

# -----------------------------
# Heading detection based on Title Case logic
# -----------------------------
def is_title_case_heading(text, known_authors=None, min_ratio=0.75, max_words=10):
    if known_authors:
        lower_text = text.lower()
        for author in known_authors:
            if author.lower() in lower_text:
                return False
    words = text.strip().split()
    if len(words) == 0 or len(words) > max_words:
        return False
    title_case_words = [w for w in words if w[0].isupper()]
    ratio = len(title_case_words) / len(words)
    return ratio >= min_ratio

# -----------------------------
# Parse body content and structure formatting
# -----------------------------
def parse_body_with_formatting(filepath, ignore_texts=None, known_authors=None):
    ignore_texts = set(ignore_texts) if ignore_texts else set()
    doc = Document(filepath)
    structured_body = []
    rels = doc.part.rels
    block_iter = iter(doc.element.body.iterchildren())

    for element in block_iter:
        tag = element.tag.lower()

        if "tbl" in tag: # Table detection
            table = next((t for t in doc.tables if t._element == element), None)
            if not table:
                continue
            table_rows = []
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_runs = []
                    for para in cell.paragraphs:
                        for run in para.runs:
                            cell_runs.append({
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline,
                                "hyperlink": None
                            })
                    row_cells.append(cell_runs)
                table_rows.append(row_cells)
            structured_body.append({
                "type": "table",
                "rows": table_rows
            })

        elif "p" in tag: # Detect and assign formatting to paragraph styles
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if not para:
                continue

            text = para.text.strip()
            if not text or text in ignore_texts:
                continue

            # Determine block type
            if is_list_paragraph(para):
                block_type = "list_item"
            elif is_title_case_heading(text, known_authors):
                block_type = "heading2"
            else:
                block_type = "paragraph"

            runs = []
            seen_hyperlink_texts = set()

            for child in para._element.iter():
                # Handle hyperlinked text
                if child.tag.endswith("}hyperlink"):
                    r_id = child.get(qn("r:id"))
                    href = rels[r_id].target_ref if r_id in rels else None
                    for node in child.iter():
                        if node.tag.endswith("}t") and node.text:
                            runs.append({
                                "text": node.text,
                                "bold": None,
                                "italic": None,
                                "underline": None,
                                "hyperlink": href
                            })
                            seen_hyperlink_texts.add(node.text)

                # Handle regular runs
                elif child.tag.endswith("}t") and child.text:
                    if child.text in seen_hyperlink_texts:
                        continue  # Skip duplicate text that was already hyperlinked
                    run_ancestor = child.getparent()
                    bold = any(e.tag.endswith("}b") for e in run_ancestor.iter())
                    italic = any(e.tag.endswith("}i") for e in run_ancestor.iter())
                    underline = any(e.tag.endswith("}u") for e in run_ancestor.iter())
                    runs.append({
                        "text": child.text,
                        "bold": bold,
                        "italic": italic,
                        "underline": underline,
                        "hyperlink": None
                    })

            structured_body.append({
                "type": block_type,
                "text": text,
                "runs": runs
            })

    return structured_body

# -----------------------------
# Main parsing function
# -----------------------------
def parse_doc(filepath, known_authors):
    doc = load_docx(filepath)
    textboxes = extract_textboxes(doc)
    title = choose_title_box(textboxes)
    structured_body = parse_body_with_formatting(filepath, ignore_texts=textboxes, known_authors=known_authors)
    author = detect_author(structured_body, known_authors)

    result = {
        "title": title,
        "structured_body": structured_body
    }
    if author:
        result["author"] = author
    return result